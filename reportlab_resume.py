#!/usr/bin/env python3
"""
Professional Resume Generator using ReportLab
Generates resumes from JSON data files with multiple format support and configurable color schemes
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, white
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import json
from pathlib import Path
import argparse


class ResumeGenerator:
    """Main resume generator class that loads data from JSON"""

    def __init__(self, data_file, config_file=None):
        self.data = self.load_resume_data(data_file)
        self.config = self.load_config(config_file)
        self.styles = self._create_styles()
        self.story = []

    def load_resume_data(self, data_file):
        """Load resume data from JSON file"""
        try:
            with open(data_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            print(f"✅ Loaded resume data from: {data_file}")
            return data
        except Exception as e:
            print(f"❌ Error loading resume data from {data_file}: {e}")
            raise

    def load_config(self, config_file):
        """Load configuration from JSON file"""
        if config_file and Path(config_file).exists():
            try:
                with open(config_file, "r", encoding="utf-8") as f:
                    config = json.load(f)
                print(f"✅ Loaded config from: {config_file}")
                return config
            except Exception as e:
                print(f"❌ Error loading config from {config_file}: {e}")

        # Return default config if no config file or error
        print("📝 Using default configuration")
        return {
            "NAME_COLOR": "#228B22",
            "TITLE_COLOR": "#B8860B",
            "SECTION_HEADER_COLOR": "#B8860B",
            "JOB_TITLE_COLOR": "#722F37",
            "ACCENT_COLOR": "#722F37",
            "COMPETENCY_HEADER_COLOR": "#228B22",
            "SUBTITLE_COLOR": "#228B22",
            "LINK_COLOR": "#B8860B",
            "DARK_TEXT_COLOR": "#333333",
            "MEDIUM_TEXT_COLOR": "#666666",
            "LIGHT_TEXT_COLOR": "#999999",
            "FONT_MAIN": "Helvetica",
            "FONT_BOLD": "Helvetica-Bold",
            "FONT_ITALIC": "Helvetica-Oblique",
            "NAME_SIZE": 24,
            "TITLE_SIZE": 14,
            "SECTION_HEADER_SIZE": 12,
            "JOB_TITLE_SIZE": 11,
            "BODY_SIZE": 9,
            "CONTACT_SIZE": 9,
            "PAGE_MARGIN": 0.6,
            "SECTION_SPACING": 0.12,
            "PARAGRAPH_SPACING": 0.06,
            "LINE_SPACING": 1.15,
            "JOB_SPACING": 6,
            "CATEGORY_SPACING": 4,
            "MAX_PAGES": 2,
            "BULLET_CHAR": "▸",
        }

    def _create_styles(self):
        """Create custom paragraph styles using configurable colors"""
        styles = getSampleStyleSheet()

        # Convert color strings to HexColor objects
        name_color = HexColor(self.config.get("NAME_COLOR", "#228B22"))
        title_color = HexColor(self.config.get("TITLE_COLOR", "#B8860B"))
        section_header_color = HexColor(
            self.config.get("SECTION_HEADER_COLOR", "#B8860B")
        )
        job_title_color = HexColor(self.config.get("JOB_TITLE_COLOR", "#722F37"))
        competency_header_color = HexColor(
            self.config.get("COMPETENCY_HEADER_COLOR", "#228B22")
        )
        subtitle_color = HexColor(self.config.get("SUBTITLE_COLOR", "#228B22"))
        dark_text_color = HexColor(self.config.get("DARK_TEXT_COLOR", "#333333"))
        medium_text_color = HexColor(self.config.get("MEDIUM_TEXT_COLOR", "#666666"))

        # Name style
        styles.add(
            ParagraphStyle(
                name="NameStyle",
                parent=styles["Heading1"],
                fontSize=self.config.get("NAME_SIZE", 24),
                textColor=name_color,
                fontName=self.config.get("FONT_BOLD", "Helvetica-Bold"),
                alignment=TA_CENTER,
                spaceAfter=4,
                spaceBefore=0,
            )
        )

        # Title style
        styles.add(
            ParagraphStyle(
                name="TitleStyle",
                parent=styles["Normal"],
                fontSize=self.config.get("TITLE_SIZE", 14),
                textColor=title_color,
                fontName=self.config.get("FONT_BOLD", "Helvetica-Bold"),
                alignment=TA_CENTER,
                spaceAfter=8,
                spaceBefore=0,
            )
        )

        # Contact style
        styles.add(
            ParagraphStyle(
                name="ContactStyle",
                parent=styles["Normal"],
                fontSize=self.config.get("CONTACT_SIZE", 9),
                textColor=medium_text_color,
                fontName=self.config.get("FONT_MAIN", "Helvetica"),
                alignment=TA_CENTER,
                spaceAfter=12,
                spaceBefore=0,
            )
        )

        # Section header style
        styles.add(
            ParagraphStyle(
                name="SectionHeader",
                parent=styles["Heading2"],
                fontSize=self.config.get("SECTION_HEADER_SIZE", 12),
                textColor=section_header_color,
                fontName=self.config.get("FONT_BOLD", "Helvetica-Bold"),
                alignment=TA_LEFT,
                spaceAfter=6,
                spaceBefore=8,
            )
        )

        # Job title style
        styles.add(
            ParagraphStyle(
                name="JobTitle",
                parent=styles["Normal"],
                fontSize=self.config.get("JOB_TITLE_SIZE", 11),
                textColor=job_title_color,
                fontName=self.config.get("FONT_BOLD", "Helvetica-Bold"),
                spaceAfter=2,
                spaceBefore=4,
            )
        )

        # Company info style
        styles.add(
            ParagraphStyle(
                name="CompanyInfo",
                parent=styles["Normal"],
                fontSize=self.config.get("BODY_SIZE", 9),
                textColor=medium_text_color,
                fontName=self.config.get("FONT_MAIN", "Helvetica"),
                spaceAfter=2,
            )
        )

        # Subtitle style
        styles.add(
            ParagraphStyle(
                name="SubtitleStyle",
                parent=styles["Normal"],
                fontSize=self.config.get("BODY_SIZE", 9),
                textColor=subtitle_color,
                fontName=self.config.get("FONT_ITALIC", "Helvetica-Oblique"),
                spaceAfter=4,
            )
        )

        # Body text style
        styles.add(
            ParagraphStyle(
                name="ResumeBodyText",
                parent=styles["Normal"],
                fontSize=self.config.get("BODY_SIZE", 9),
                textColor=dark_text_color,
                fontName=self.config.get("FONT_MAIN", "Helvetica"),
                alignment=TA_JUSTIFY,
                spaceAfter=2,
                leading=self.config.get("BODY_SIZE", 9)
                * self.config.get("LINE_SPACING", 1.15),
            )
        )

        # Bullet style
        styles.add(
            ParagraphStyle(
                name="ResumeBulletStyle",
                parent=styles["Normal"],
                fontSize=self.config.get("BODY_SIZE", 9),
                textColor=dark_text_color,
                fontName=self.config.get("FONT_MAIN", "Helvetica"),
                leftIndent=12,
                bulletIndent=0,
                spaceAfter=1.5,
                leading=self.config.get("BODY_SIZE", 9)
                * self.config.get("LINE_SPACING", 1.15),
            )
        )

        # Competency header style
        styles.add(
            ParagraphStyle(
                name="CompetencyHeader",
                parent=styles["Normal"],
                fontSize=self.config.get("BODY_SIZE", 9) + 1,
                textColor=competency_header_color,
                fontName=self.config.get("FONT_BOLD", "Helvetica-Bold"),
                spaceAfter=3,
                spaceBefore=4,
            )
        )

        return styles

    def _add_header(self):
        """Add the header section with name, title, and contact info"""
        personal_info = self.data.get("personal_info", {})

        # Name
        name_para = Paragraph(
            personal_info.get("name", "NAME"), self.styles["NameStyle"]
        )
        self.story.append(name_para)

        # Title
        title_para = Paragraph(
            personal_info.get("title", "Professional Title"), self.styles["TitleStyle"]
        )
        self.story.append(title_para)

        # Contact info - using configurable link color
        link_color = self.config.get("LINK_COLOR", "#B8860B")
        contact_text = f"""
        <b>{personal_info.get('phone', '')} | {personal_info.get('email', '')}</b><br/>
        <a href="{personal_info.get('website', '')}" color="{link_color}">{personal_info.get('website', '')}</a> |
        <a href="{personal_info.get('linkedin', '')}" color="{link_color}">{personal_info.get('linkedin', '')}</a>
        """
        contact_para = Paragraph(contact_text, self.styles["ContactStyle"])
        self.story.append(contact_para)

        # Header separator
        self.story.append(Spacer(1, 4))

    def _add_section_header(self, title):
        """Add a section header with underline"""
        header_text = f"<u>{title.upper()}</u>"
        header_para = Paragraph(header_text, self.styles["SectionHeader"])
        self.story.append(header_para)

    def _add_summary(self):
        """Add professional summary section"""
        summary = self.data.get("summary", "")
        if summary:
            self._add_section_header("Professional Summary")
            summary_para = Paragraph(summary, self.styles["ResumeBodyText"])
            self.story.append(summary_para)
            self.story.append(
                Spacer(1, self.config.get("SECTION_SPACING", 0.12) * inch)
            )

    def _add_competencies(self):
        """Add core competencies section"""
        competencies = self.data.get("competencies", {})
        if competencies:
            self._add_section_header("Core Competencies")

            for category, skills in competencies.items():
                # Category header
                cat_para = Paragraph(category, self.styles["CompetencyHeader"])
                self.story.append(cat_para)

                # Skills list
                if isinstance(skills, list):
                    if any(
                        "Programming:" in skill or "Data Platforms:" in skill
                        for skill in skills
                    ):
                        # Technical skills with better formatting
                        for skill in skills:
                            skill_para = Paragraph(skill, self.styles["ResumeBodyText"])
                            self.story.append(skill_para)
                    else:
                        # Other competencies as a flowing paragraph
                        skills_text = " • ".join(skills)
                        skills_para = Paragraph(
                            skills_text, self.styles["ResumeBodyText"]
                        )
                        self.story.append(skills_para)
                else:
                    # Handle string values
                    skills_para = Paragraph(str(skills), self.styles["ResumeBodyText"])
                    self.story.append(skills_para)

                self.story.append(Spacer(1, self.config.get("CATEGORY_SPACING", 4)))

            self.story.append(
                Spacer(1, self.config.get("SECTION_SPACING", 0.12) * inch)
            )

    def _add_experience(self):
        """Add professional experience section"""
        experience = self.data.get("experience", [])
        if experience:
            self._add_section_header("Professional Experience")

            for i, job in enumerate(experience):
                job_content = []

                # Job title
                title_para = Paragraph(job.get("title", ""), self.styles["JobTitle"])
                job_content.append(title_para)

                # Company and dates
                company_para = Paragraph(
                    f"{job.get('company', '')} | {job.get('dates', '')}",
                    self.styles["CompanyInfo"],
                )
                job_content.append(company_para)

                # Subtitle
                if job.get("subtitle"):
                    subtitle_para = Paragraph(
                        job["subtitle"], self.styles["SubtitleStyle"]
                    )
                    job_content.append(subtitle_para)

                # Responsibilities
                responsibilities = job.get("responsibilities", [])
                for responsibility in responsibilities:
                    bullet_text = (
                        f"{self.config.get('BULLET_CHAR', '▸')} {responsibility}"
                    )
                    bullet_para = Paragraph(
                        bullet_text, self.styles["ResumeBulletStyle"]
                    )
                    job_content.append(bullet_para)

                # Keep job together on same page if possible
                job_group = KeepTogether(job_content)
                self.story.append(job_group)

                # Add spacing between jobs
                if i < len(experience) - 1:
                    self.story.append(Spacer(1, self.config.get("JOB_SPACING", 6)))

            # Add LinkedIn reference
            personal_info = self.data.get("personal_info", {})
            linkedin_url = personal_info.get("linkedin", "")
            if linkedin_url:
                link_color = self.config.get("LINK_COLOR", "#B8860B")
                linkedin_text = f'<i>Additional experience and project details available on <a href="{linkedin_url}" color="{link_color}">LinkedIn</a></i>'
                linkedin_para = Paragraph(linkedin_text, self.styles["ResumeBodyText"])
                self.story.append(Spacer(1, 4))
                self.story.append(linkedin_para)
                self.story.append(
                    Spacer(1, self.config.get("SECTION_SPACING", 0.12) * inch)
                )

    def _add_achievements(self):
        """Add key achievements section"""
        achievements = self.data.get("achievements", {})
        if achievements:
            self._add_section_header("Key Achievements and Impact")

            for category, achievement_list in achievements.items():
                # Category header
                cat_para = Paragraph(category, self.styles["CompetencyHeader"])
                self.story.append(cat_para)

                # Achievements list
                if isinstance(achievement_list, list):
                    for achievement in achievement_list:
                        bullet_text = f"✓ {achievement}"
                        bullet_para = Paragraph(
                            bullet_text, self.styles["ResumeBulletStyle"]
                        )
                        self.story.append(bullet_para)

                self.story.append(Spacer(1, 4))

    def generate_pdf(self, filename="resume.pdf"):
        """Generate the complete PDF resume"""
        # Create document
        doc = SimpleDocTemplate(
            filename,
            pagesize=letter,
            rightMargin=self.config.get("PAGE_MARGIN", 0.6) * inch,
            leftMargin=self.config.get("PAGE_MARGIN", 0.6) * inch,
            topMargin=self.config.get("PAGE_MARGIN", 0.6) * inch,
            bottomMargin=self.config.get("PAGE_MARGIN", 0.6) * inch,
        )

        # Build content
        self._add_header()
        self._add_summary()
        self._add_competencies()
        self._add_experience()
        self._add_achievements()

        # Generate PDF
        doc.build(self.story)
        print(f"✅ PDF generated: {filename}")
        return filename

    def generate_docx(self, filename="resume.docx"):
        """Generate Word document version"""
        try:
            from docx import Document
            from docx.shared import Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("⚠️  python-docx not installed. Install with: pip install python-docx")
            return None

        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            margin = self.config.get("PAGE_MARGIN", 0.6)
            section.top_margin = Inches(margin)
            section.bottom_margin = Inches(margin)
            section.left_margin = Inches(margin)
            section.right_margin = Inches(margin)

        personal_info = self.data.get("personal_info", {})

        # Helper function to convert hex to RGB
        def hex_to_rgb(hex_color):
            hex_color = hex_color.lstrip("#")
            return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))

        # Header
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal_info.get("name", "NAME"))
        name_run.font.size = Inches(0.3)  # Approximate 24pt
        name_run.font.bold = True
        name_rgb = hex_to_rgb(self.config.get("NAME_COLOR", "#228B22"))
        name_run.font.color.rgb = RGBColor(*name_rgb)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(personal_info.get("title", "Professional Title"))
        title_run.font.size = Inches(0.17)  # Approximate 14pt
        title_run.font.bold = True
        title_rgb = hex_to_rgb(self.config.get("TITLE_COLOR", "#B8860B"))
        title_run.font.color.rgb = RGBColor(*title_rgb)

        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = f"{personal_info.get('phone', '')} | {personal_info.get('email', '')}\n{personal_info.get('website', '')} | {personal_info.get('linkedin', '')}"
        contact_para.add_run(contact_text)

        # Summary
        summary = self.data.get("summary", "")
        if summary:
            self._add_docx_section(doc, "PROFESSIONAL SUMMARY", summary)

        # Competencies
        competencies = self.data.get("competencies", {})
        if competencies:
            comp_heading = doc.add_heading("CORE COMPETENCIES", level=2)
            section_rgb = hex_to_rgb(self.config.get("SECTION_HEADER_COLOR", "#B8860B"))
            comp_heading.runs[0].font.color.rgb = RGBColor(*section_rgb)

            for category, skills in competencies.items():
                cat_para = doc.add_paragraph()
                cat_run = cat_para.add_run(category)
                cat_run.font.bold = True
                comp_rgb = hex_to_rgb(
                    self.config.get("COMPETENCY_HEADER_COLOR", "#228B22")
                )
                cat_run.font.color.rgb = RGBColor(*comp_rgb)

                if isinstance(skills, list):
                    if any(
                        "Programming:" in skill or "Data Platforms:" in skill
                        for skill in skills
                    ):
                        for skill in skills:
                            doc.add_paragraph(skill, style="List Bullet")
                    else:
                        skills_text = " • ".join(skills)
                        doc.add_paragraph(skills_text)

        # Experience
        experience = self.data.get("experience", [])
        if experience:
            exp_heading = doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
            section_rgb = hex_to_rgb(self.config.get("SECTION_HEADER_COLOR", "#B8860B"))
            exp_heading.runs[0].font.color.rgb = RGBColor(*section_rgb)

            for job in experience:
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(job.get("title", ""))
                job_run.font.bold = True
                job_rgb = hex_to_rgb(self.config.get("JOB_TITLE_COLOR", "#722F37"))
                job_run.font.color.rgb = RGBColor(*job_rgb)

                company_para = doc.add_paragraph(
                    f"{job.get('company', '')} | {job.get('dates', '')}"
                )

                if job.get("subtitle"):
                    subtitle_para = doc.add_paragraph()
                    subtitle_run = subtitle_para.add_run(job["subtitle"])
                    subtitle_run.font.italic = True
                    subtitle_rgb = hex_to_rgb(
                        self.config.get("SUBTITLE_COLOR", "#228B22")
                    )
                    subtitle_run.font.color.rgb = RGBColor(*subtitle_rgb)

                responsibilities = job.get("responsibilities", [])
                for responsibility in responsibilities:
                    doc.add_paragraph(f"▸ {responsibility}", style="List Bullet")

        # Achievements
        achievements = self.data.get("achievements", {})
        if achievements:
            ach_heading = doc.add_heading("KEY ACHIEVEMENTS AND IMPACT", level=2)
            section_rgb = hex_to_rgb(self.config.get("SECTION_HEADER_COLOR", "#B8860B"))
            ach_heading.runs[0].font.color.rgb = RGBColor(*section_rgb)

            for category, achievement_list in achievements.items():
                cat_para = doc.add_paragraph()
                cat_run = cat_para.add_run(category)
                cat_run.font.bold = True
                comp_rgb = hex_to_rgb(
                    self.config.get("COMPETENCY_HEADER_COLOR", "#228B22")
                )
                cat_run.font.color.rgb = RGBColor(*comp_rgb)

                if isinstance(achievement_list, list):
                    for achievement in achievement_list:
                        doc.add_paragraph(f"✓ {achievement}", style="List Bullet")

        doc.save(filename)
        print(f"✅ DOCX generated: {filename}")
        return filename

    def _add_docx_section(self, doc, title, content):
        """Helper to add a section to Word doc"""
        from docx.shared import RGBColor

        def hex_to_rgb(hex_color):
            hex_color = hex_color.lstrip("#")
            return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))

        heading = doc.add_heading(title, level=2)
        section_rgb = hex_to_rgb(self.config.get("SECTION_HEADER_COLOR", "#B8860B"))
        heading.runs[0].font.color.rgb = RGBColor(*section_rgb)
        doc.add_paragraph(content)

    def generate_rtf(self, filename="resume.rtf"):
        """Generate RTF document"""
        personal_info = self.data.get("personal_info", {})

        # Helper function to convert hex to RGB for RTF
        def hex_to_rtf_rgb(hex_color):
            hex_color = hex_color.lstrip("#")
            r, g, b = tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))
            return f"\\red{r}\\green{g}\\blue{b}"

        # Build color table
        name_rgb = hex_to_rtf_rgb(self.config.get("NAME_COLOR", "#228B22"))
        title_rgb = hex_to_rtf_rgb(self.config.get("TITLE_COLOR", "#B8860B"))
        section_rgb = hex_to_rtf_rgb(self.config.get("SECTION_HEADER_COLOR", "#B8860B"))
        job_rgb = hex_to_rtf_rgb(self.config.get("JOB_TITLE_COLOR", "#722F37"))

        rtf_content = f"""{{\\rtf1\\ansi\\deff0
{{\\fonttbl{{\\f0 Times New Roman;}}}}
{{\\colortbl;\\red0\\green0\\blue0;{name_rgb};{title_rgb};{section_rgb};{job_rgb};}}
\\f0\\fs24
{{\\qc\\cf2\\b\\fs36 {personal_info.get('name', 'NAME')}\\par}}
{{\\qc\\cf3\\b\\fs20 {personal_info.get('title', 'Professional Title')}\\par}}
{{\\qc {personal_info.get('phone', '')} | {personal_info.get('email', '')}\\par}}
{{\\qc {personal_info.get('website', '')} | {personal_info.get('linkedin', '')}\\par}}
\\par
"""

        # Summary
        summary = self.data.get("summary", "")
        if summary:
            rtf_content += (
                r"{\cf4\b\ul PROFESSIONAL SUMMARY\par}" + summary + r"\par\par"
            )

        # Competencies
        competencies = self.data.get("competencies", {})
        if competencies:
            rtf_content += r"{\cf4\b\ul CORE COMPETENCIES\par}"
            for category, skills in competencies.items():
                rtf_content += r"{\cf2\b " + category + r"\par}"
                if isinstance(skills, list):
                    if any(
                        "Programming:" in skill or "Data Platforms:" in skill
                        for skill in skills
                    ):
                        for skill in skills:
                            rtf_content += skill + r"\par"
                    else:
                        rtf_content += " • ".join(skills) + r"\par"
                rtf_content += r"\par"

        # Experience
        experience = self.data.get("experience", [])
        if experience:
            rtf_content += r"{\cf4\b\ul PROFESSIONAL EXPERIENCE\par}"
            for job in experience:
                rtf_content += r"{\cf5\b " + job.get("title", "") + r"\par}"
                rtf_content += (
                    job.get("company", "") + " | " + job.get("dates", "") + r"\par"
                )
                if job.get("subtitle"):
                    rtf_content += r"{\cf2\i " + job["subtitle"] + r"\par}"
                responsibilities = job.get("responsibilities", [])
                for resp in responsibilities:
                    rtf_content += "▸ " + resp + r"\par"
                rtf_content += r"\par"

        # Achievements
        achievements = self.data.get("achievements", {})
        if achievements:
            rtf_content += r"{\cf4\b\ul KEY ACHIEVEMENTS AND IMPACT\par}"
            for category, achievement_list in achievements.items():
                rtf_content += r"{\cf2\b " + category + r"\par}"
                if isinstance(achievement_list, list):
                    for achievement in achievement_list:
                        rtf_content += "✓ " + achievement + r"\par"
                rtf_content += r"\par"

        rtf_content += "}"

        with open(filename, "w", encoding="utf-8") as f:
            f.write(rtf_content)

        print(f"✅ RTF generated: {filename}")
        return filename


def main():
    """Main function to generate resume"""
    parser = argparse.ArgumentParser(
        description="Generate professional resume from JSON data"
    )
    parser.add_argument(
        "--format",
        choices=["pdf", "docx", "rtf", "all"],
        default="pdf",
        help="Output format(s) to generate",
    )
    parser.add_argument(
        "--basename",
        required=True,
        help="Base name for input directory and output files",
    )
    parser.add_argument(
        "--input-dir", default="inputs", help="Input directory containing resume data"
    )
    parser.add_argument(
        "--output-dir", default="outputs", help="Output directory for generated files"
    )

    args = parser.parse_args()

    # Construct paths
    input_base = Path(args.input_dir) / args.basename
    data_file = input_base / "resume_data.json"
    config_file = input_base / "config.json"

    # Check if input files exist
    if not data_file.exists():
        print(f"❌ Resume data file not found: {data_file}")
        print(f"Available directories in {args.input_dir}:")
        input_dir = Path(args.input_dir)
        if input_dir.exists():
            for item in input_dir.iterdir():
                if item.is_dir():
                    print(f"   • {item.name}")
        return 1

    # Create output directory structure
    output_base = Path(args.output_dir) / args.basename
    formats_to_generate = (
        ["pdf", "docx", "rtf"] if args.format == "all" else [args.format]
    )

    for fmt in formats_to_generate:
        (output_base / fmt).mkdir(parents=True, exist_ok=True)

    try:
        # Generate resume
        generator = ResumeGenerator(data_file, config_file)
        generated_files = []

        for fmt in formats_to_generate:
            output_dir = output_base / fmt
            filename = output_dir / f"{args.basename}.{fmt}"

            if fmt == "pdf":
                result = generator.generate_pdf(str(filename))
            elif fmt == "docx":
                result = generator.generate_docx(str(filename))
            elif fmt == "rtf":
                result = generator.generate_rtf(str(filename))

            if result:
                generated_files.append(result)

        print(f"\n🎉 Resume generation complete!")
        for file in generated_files:
            print(f"📄 Generated: {file}")

        return 0

    except Exception as e:
        print(f"❌ Error generating resume: {e}")
        return 1


if __name__ == "__main__":
    exit(main())
