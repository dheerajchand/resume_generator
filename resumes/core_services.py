#!/usr/bin/env python3
"""
Core Resume Generation Services - MAIN FORMAT
Consolidated functionality for resume generation across all formats

This is the MAIN FORMAT with optimized spacing and space-efficient design.
For the minimal/backup format, see core_services_minimal.py
"""

import json
import os
from pathlib import Path
from typing import Dict, List, Optional, Any
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, white
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether, Table, TableStyle, PageTemplate, BaseDocTemplate
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.platypus.frames import Frame
from reportlab.platypus.doctemplate import PageTemplate
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from resume_generator_django.resume_generator.constants import (
    SPACE_BASE, SPACE_BETWEEN_SECTIONS, SPACE_BETWEEN_JOB_UNITS, 
    SPACE_BETWEEN_JOB_COMPONENTS, SPACE_HEADER_TO_CONTENT, SPACE_SUBHEADER_TO_BULLETS, SPACE_HEADER_TOP,
    SPACE_HEADER_BAR_TO_CONTENT, SPACE_HEADER_HEIGHT, SPACE_FOOTER_HEIGHT, SPACE_HEADER_TO_MAIN_BODY,
    FONT_SIZE_SECTION_HEADER, FONT_SIZE_COMPANY, FONT_SIZE_MAIN_COMPETENCY,
    FONT_SIZE_JOB_TITLE, FONT_SIZE_BODY, FONT_SIZE_SUB_COMPETENCY,
    FONT_SIZE_BULLET_POINT, FONT_SIZE_COMPETENCY_DETAIL, FONT_SIZE_FOOTER,
    COLOR_MAPPINGS, PARAGRAPH_SPACING, SECTION_ORDER,
    MARGIN_LEFT, MARGIN_RIGHT, MARGIN_TOP, MARGIN_BOTTOM, PAGE_HEIGHT,
    HEADER_LEFT_X, HEADER_TOP_Y, HEADER_RIGHT_X, HEADER_FOOTER_Y,
    HEADER_PHONE_OFFSET, HEADER_GITHUB_OFFSET, HEADER_GITHUB_LINK_OFFSET,
    HEADER_NAME_OFFSET, HEADER_LOCATION_OFFSET, HEADER_LOCATION_LINK_OFFSET,
    FOOTER_LINK_OFFSET, FOOTER_BAR_POSITION, SPACE_MULTIPLIER_TINY, SPACE_MULTIPLIER_MINIMAL, SPACE_MULTIPLIER_SMALL, 
    SPACE_MULTIPLIER_MEDIUM, SPACE_MULTIPLIER_LARGE, FONT_SIZE_8, FONT_SIZE_9, 
    FONT_SIZE_10, FONT_SIZE_11, FONT_SIZE_12, FONT_SIZE_14, BAR_WIDTH, BAR_HEIGHT,
    BAR_LINE_WIDTH, BAR_LINE_WIDTH_FOOTER, PAGE_CONTENT_WIDTH, PAGE_LEFT_MARGIN,
    PAGE_RIGHT_MARGIN, HEADER_FIRST_PHONE_Y, HEADER_FIRST_GITHUB_Y, 
    HEADER_FIRST_NAME_Y, HEADER_FIRST_LOCATION_Y, HEADER_RECURRING_NAME_Y,
    HEADER_RECURRING_EMAIL_Y, HEADER_RECURRING_PHONE_Y, HEADER_RECURRING_GITHUB_Y,
    FOOTER_Y, FONT_THEMES, FONT_ROLES, FONT_SIZE_THEMES,
    get_spacing_constant, get_font_size, get_color_role, get_theme_font, get_theme_font_size
)




class ResumeGenerator:
    """Core resume generator supporting all formats"""
    
    def __init__(self, data_file: str, config_file: Optional[str] = None, color_scheme: str = 'default_professional', length_variant: str = 'long', output_type: str = 'ats'):
        self.data = self._load_json(data_file)
        self.config = self._load_json(config_file) if config_file else {}
        self.color_scheme = color_scheme
        self.length_variant = length_variant
        self.output_type = output_type
        self.styles = self._create_styles()
        
        # Spacing system constants (imported from settings)
        self.SPACE_BASE = SPACE_BASE
        self.SPACE_BETWEEN_SECTIONS = SPACE_BETWEEN_SECTIONS
        self.SPACE_BETWEEN_JOB_UNITS = SPACE_BETWEEN_JOB_UNITS
        self.SPACE_BETWEEN_JOB_COMPONENTS = SPACE_BETWEEN_JOB_COMPONENTS
        self.SPACE_HEADER_TO_CONTENT = SPACE_HEADER_TO_CONTENT
        self.SPACE_SUBHEADER_TO_BULLETS = SPACE_SUBHEADER_TO_BULLETS
    
    def _load_json(self, file_path: str) -> Dict[str, Any]:
        """Load JSON data from file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            raise Exception(f"Error loading {file_path}: {e}")
    
    def _create_styles(self) -> Dict[str, ParagraphStyle]:
        """Create paragraph styles based on config with theme-specific fonts"""
        styles = getSampleStyleSheet()
        
        # Get colors from config (color schemes have colors directly, not wrapped in "colors")
        colors = self.config if self.config else {}
        
        # Get color scheme name for font theme selection
        color_scheme = getattr(self, 'color_scheme', 'default_professional')
        
        # Helper function to get theme-specific font with proper bold/italic handling
        def get_font(role, bold=False, italic=False):
            base_font = get_theme_font(color_scheme, role)
            
            # If the theme already specifies bold/italic, use it directly
            if '-Bold' in base_font or '-Oblique' in base_font:
                return base_font
            
            # Handle font variations for ReportLab
            if bold and italic:
                return f"{base_font}-BoldOblique"
            elif bold:
                return f"{base_font}-Bold"
            elif italic:
                return f"{base_font}-Oblique"
            else:
                return base_font
        
        custom_styles = {
            "Name": ParagraphStyle(
                "CustomName",
                parent=styles["Heading1"],
                fontSize=28,  # Bigger name - keep as special case, consistent across all themes
                textColor=HexColor(colors.get("NAME_COLOR", "#2C3E50")),
                alignment=TA_RIGHT,
                spaceAfter=get_spacing_constant('base') * 1,
                fontName=get_font('name', bold=True),
            ),
            "Title": ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading2"],
                fontSize=get_font_size('section_header'),
                textColor=HexColor(colors.get("TITLE_COLOR", "#34495E")),
                alignment=TA_CENTER,
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_LARGE,
                fontName=get_font('primary', bold=True),
            ),
            "Subtitle": ParagraphStyle(
                "CustomSubtitle",
                parent=styles["Normal"],
                fontSize=get_font_size('bullet_point'),
                textColor=HexColor(colors.get("TITLE_COLOR", "#7F8C8D")),
                alignment=TA_CENTER,
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_LARGE,
                fontName=get_font('secondary'),
            ),
            # DESIGN SYSTEM: Consistent spacing scale (0.25, 0.5, 1, 2, 4 units)
            # Typography hierarchy: 8, 9, 10, 11, 12, 14pt
            # Color hierarchy: primary, secondary, accent, muted
            
            "SectionHeader": ParagraphStyle(
                "CustomSectionHeader",
                parent=styles["Heading2"],
                fontSize=get_font_size('section_header'),
                textColor=HexColor(colors.get("SECTION_HEADER_COLOR", "#2C3E50")),
                spaceAfter=get_spacing_constant('header_to_content'),  # Minimal gap to content
                spaceBefore=get_spacing_constant('base') * SPACE_MULTIPLIER_MEDIUM,     # Medium units spacing - reduced whitespace between main sections
                fontName=get_font('primary', bold=True),
            ),
            "JobTitle": ParagraphStyle(
                "CustomJobTitle",
                parent=styles["Normal"],
                fontSize=get_font_size('job_title'),
                textColor=HexColor(colors.get("JOB_TITLE_COLOR", "#666666")),  # Muted color
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_SMALL,   # Small spacing - increased gap between job title and subtitle
                spaceBefore=get_spacing_constant('base') * SPACE_MULTIPLIER_SMALL,  # Small spacing
                fontName=get_font('accent'),
            ),
            "Company": ParagraphStyle(
                "CustomCompany",
                parent=styles["Heading3"],
                fontSize=get_font_size('company'),
                textColor=HexColor(colors.get("COMPANY_COLOR", "#2C3E50")),  # Primary color
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_MINIMAL,    # Minimal spacing after company
                spaceBefore=get_spacing_constant('base') * SPACE_MULTIPLIER_MINIMAL,   # Minimal spacing before company
                fontName=get_font('primary', bold=True),
            ),
            "Body": ParagraphStyle(
                "CustomBody",
                parent=styles["Normal"],
                fontSize=get_font_size('body'),
                textColor=HexColor(colors.get("DARK_TEXT_COLOR", "#2C3E50")),
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_MEDIUM,      # Medium unit spacing
                fontName=get_font('secondary'),
            ),
            "BulletPoint": ParagraphStyle(
                "CustomBulletPoint",
                parent=styles["Normal"],
                fontSize=get_font_size('bullet_point'),
                textColor=HexColor(colors.get("MEDIUM_TEXT_COLOR", "#666666")),
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_SMALL,    # Small unit spacing - tighter bullet spacing
                leftIndent=12,
                fontName=get_font('secondary'),
            ),
            "MainCompetency": ParagraphStyle(
                "CustomMainCompetency",
                parent=styles["Normal"],
                fontSize=get_font_size('main_competency'),
                textColor=HexColor(colors.get("COMPETENCY_HEADER_COLOR", "#2C3E50")),
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_MINIMAL,    # Minimal unit spacing - minimal gap to bullets
                spaceBefore=get_spacing_constant('base') * SPACE_MULTIPLIER_MINIMAL,   # Minimal unit spacing - minimal gap from previous
                fontName=get_font('primary', bold=True),
            ),
            "SubCompetency": ParagraphStyle(
                "CustomSubCompetency",
                parent=styles["Normal"],
                fontSize=get_font_size('sub_competency'),
                textColor=HexColor(colors.get("ACCENT_COLOR", "#4682B4")),
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_MEDIUM,    # Medium unit spacing
                leftIndent=12,
                fontName=get_font('accent', bold=True),
            ),
            "CompetencyDetail": ParagraphStyle(
                "CustomCompetencyDetail",
                parent=styles["Normal"],
                fontSize=get_font_size('competency_detail'),
                textColor=HexColor(colors.get("DARK_TEXT_COLOR", "#2C3E50")),
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_MEDIUM,    # Medium unit spacing
                leftIndent=0,
                fontName=get_font('secondary'),
                allowWidows=1,
                allowOrphans=1,
            ),
            "Contact": ParagraphStyle(
                "CustomContact",
                parent=styles["Normal"],
                fontSize=get_font_size('body'),
                textColor=HexColor(colors.get("ACCENT_COLOR", "#4682B4")),  # Use accent color for header contact
                alignment=TA_RIGHT,
                spaceAfter=get_spacing_constant('base') * 1,
                fontName=get_font('secondary'),
            ),
            "ContactStacked": ParagraphStyle(
                "CustomContactStacked",
                parent=styles["Normal"],
                fontSize=get_font_size('body'),
                textColor=HexColor(colors.get("ACCENT_COLOR", "#4682B4")),
                alignment=TA_RIGHT,
                spaceAfter=get_spacing_constant('base') * SPACE_MULTIPLIER_MEDIUM,
                fontName=get_font('secondary'),
            ),
        }
        
        return custom_styles
    
    def _get_sections(self):
        """Define and render all resume sections in order"""
        sections = []
        
        # Professional Summary
        summary = self.data.get("summary", "")
        if summary:
            sections.append({
                "name": "PROFESSIONAL SUMMARY",
                "content": self._create_keep_together_section([Paragraph(summary, self.styles["Body"])])
            })
        
        # Key Achievements and Impact (abbreviated like Deepak)
        achievements = self.data.get("achievements", {})
        if achievements:
            achievement_content = []
            # Single line with bullet separators (abbreviated format)
            all_achievements = []
            for category, achievement_list in achievements.items():
                if isinstance(achievement_list, list) and achievement_list:
                    # Add all achievements from the list
                    all_achievements.extend(achievement_list)
            
            # Single line format like Deepak with bullet separators
            achievement_text = " • ".join(all_achievements)
            achievement_content.append(Paragraph(achievement_text, self.styles["CompetencyDetail"]))
            
            sections.append({
                "name": "KEY ACHIEVEMENTS AND IMPACT",
                "content": self._create_keep_together_section(achievement_content)
            })
        
        # CORE COMPETENCIES - Use actual categories from data
        competencies_data = self.data.get("competencies", {})
        if competencies_data:
            core_competencies = list(competencies_data.keys())
            
            # Single line with bullet separators (like Deepak)
            competency_text = " • ".join(core_competencies)
            competency_content = [Paragraph(competency_text, self.styles["CompetencyDetail"])]
        else:
            competency_content = []
        
        sections.append({
            "name": "CORE COMPETENCIES",
            "content": self._create_keep_together_section(competency_content)
        })
        
        # Professional Experience
        experience = self.data.get("experience", [])
        if experience:
            experience_content = []
            for job in experience:
                job_title = job.get("title", "")
                company = job.get("company", "")
                location = job.get("location", "")
                dates = job.get("dates", "")
                
                # Company, job title, and dates on one line
                company_line = company
                if job_title:
                    company_line += f" | {job_title}"
                if location:
                    company_line += f" - {location}"
                if dates:
                    company_line += f" {dates}"
                
                # Intelligent job unit splitting logic
                job_unit = []
                job_unit.append(Paragraph(company_line, self.styles["Company"]))
                
                if job.get("subtitle"):
                    job_unit.append(Spacer(1, self.SPACE_BETWEEN_JOB_COMPONENTS))  # Spacing between company and subtitle
                    job_unit.append(Paragraph(job["subtitle"], self.styles["SubCompetency"]))
                
                responsibilities = job.get("responsibilities", [])
                
                # No additional spacing - JobTitle.spaceAfter already provides the correct spacing
                
                # If job has many responsibilities, allow splitting but keep header together
                if len(responsibilities) > 4:  # Increased threshold to 4
                    # Keep company + title + subtitle + first bullet together
                    header_unit = job_unit.copy()
                    # Add first bullet to header to ensure at least one stays with header
                    header_unit.append(Paragraph(f"• {responsibilities[0]}", self.styles["BulletPoint"]))
                    experience_content.append(KeepTogether(header_unit))
                    
                    # Add remaining bullets (can split across pages)
                    for resp in responsibilities[1:]:
                        experience_content.append(Paragraph(f"• {resp}", self.styles["BulletPoint"]))
                elif len(responsibilities) > 1:  # For jobs with 2-4 responsibilities
                    # Keep header + at least one bullet together
                    header_unit = job_unit.copy()
                    # Add first bullet to header to ensure at least one stays with header
                    header_unit.append(Paragraph(f"• {responsibilities[0]}", self.styles["BulletPoint"]))
                    experience_content.append(KeepTogether(header_unit))
                    
                    # Add remaining bullets (can split across pages)
                    for resp in responsibilities[1:]:
                        experience_content.append(Paragraph(f"• {resp}", self.styles["BulletPoint"]))
                else:
                    # For shorter jobs, keep everything together
                    for resp in responsibilities:
                        job_unit.append(Paragraph(f"• {resp}", self.styles["BulletPoint"]))
                    experience_content.append(KeepTogether(job_unit))
                
                # Add spacer between job units (but not after the last job)
                if job != experience[-1]:  # Only add spacer if not the last job
                    experience_content.append(Spacer(1, self.SPACE_BETWEEN_JOB_UNITS))
            
            sections.append({
                "name": "PROFESSIONAL EXPERIENCE",
                "content": experience_content
            })
        
        # Key Projects - Keep header with first project (minimum 4 lines)
        projects = self.data.get("projects", [])
        if projects:
            project_content = []
            
            # For the first project, include it in a KeepTogether with the section header
            # to ensure the header doesn't appear alone at the bottom of a page
            first_project = projects[0]
            first_project_name = first_project.get("name", "")
            first_dates = first_project.get("dates", "")
            first_description = first_project.get("description", "")
            first_technologies = first_project.get("technologies", [])
            first_impact = first_project.get("impact", "")
            
            # Build first project unit with header
            first_project_unit = []
            
            # Add section header as part of first project unit
            first_project_unit.append(Paragraph("KEY PROJECTS", self.styles["SectionHeader"]))
            
            # Add first project content
            first_title_line = first_project_name
            if first_dates:
                first_title_line += f" ({first_dates})"
            first_project_unit.append(Paragraph(first_title_line, self.styles["SubCompetency"]))
            
            if first_description:
                # Add "About:" with different color formatting
                colors = self.config if self.config else {}
                medium_color = colors.get("MEDIUM_TEXT_COLOR", "#666666")
                about_text = f'<font color="{medium_color}"><b>About:</b></font> {first_description}'
                first_project_unit.append(Paragraph(about_text, self.styles["CompetencyDetail"]))
            
            if first_technologies:
                # Add "Technologies:" with color formatting
                colors = self.config if self.config else {}
                tech_color = colors.get("COMPETENCY_HEADER_COLOR", "#2C3E50")
                first_tech_text = f'<font color="{tech_color}"><b>Technologies:</b></font> {", ".join(first_technologies)}'
                first_project_unit.append(Paragraph(first_tech_text, self.styles["CompetencyDetail"]))
            
            if first_impact:
                # Add "Impact:" with color formatting
                colors = self.config if self.config else {}
                accent_color = colors.get("ACCENT_COLOR", "#4682B4")
                impact_text = f'<font color="{accent_color}"><b>Impact:</b></font> {first_impact}'
                first_project_unit.append(Paragraph(impact_text, self.styles["CompetencyDetail"]))
            
            # Keep header with first project together (ensures 4+ lines)
            project_content.append(KeepTogether(first_project_unit))
            
            # Handle remaining projects normally
            for project in projects[1:]:
                project_name = project.get("name", "")
                dates = project.get("dates", "")
                description = project.get("description", "")
                technologies = project.get("technologies", [])
                impact = project.get("impact", "")
                
                # Add spacing before each subsequent project
                project_content.append(Spacer(1, self.SPACE_BETWEEN_JOB_COMPONENTS))
                
                # Build individual project unit
                project_unit = []
                
                title_line = project_name
                if dates:
                    title_line += f" ({dates})"
                project_unit.append(Paragraph(title_line, self.styles["SubCompetency"]))
                
                if description:
                    # Add "About:" with different color formatting
                    colors = self.config if self.config else {}
                    medium_color = colors.get("MEDIUM_TEXT_COLOR", "#666666")
                    about_text = f'<font color="{medium_color}"><b>About:</b></font> {description}'
                    project_unit.append(Paragraph(about_text, self.styles["CompetencyDetail"]))
                
                if technologies:
                    # Add "Technologies:" with color formatting
                    colors = self.config if self.config else {}
                    tech_color = colors.get("COMPETENCY_HEADER_COLOR", "#2C3E50")
                    tech_text = f'<font color="{tech_color}"><b>Technologies:</b></font> {", ".join(technologies)}'
                    project_unit.append(Paragraph(tech_text, self.styles["CompetencyDetail"]))
                
                if impact:
                    # Add "Impact:" with color formatting
                    colors = self.config if self.config else {}
                    accent_color = colors.get("ACCENT_COLOR", "#4682B4")
                    impact_text = f'<font color="{accent_color}"><b>Impact:</b></font> {impact}'
                    project_unit.append(Paragraph(impact_text, self.styles["CompetencyDetail"]))
                
                # Keep each individual project together
                project_content.append(KeepTogether(project_unit))
            
            # Add section without separate header (header is included in first project)
            sections.append({
                "name": "", 
                "content": project_content
            })
        
        # Education - Individual education KeepTogether logic
        education = self.data.get("education", [])
        if education:
            education_content = []
            for edu in education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                location = edu.get("location", "")
                dates = edu.get("dates", "")
                gpa = edu.get("gpa", "")
                honors = edu.get("honors", "")
                
                # Build individual education unit
                edu_unit = []
                
                title_line = degree
                if institution:
                    title_line += f" - {institution}"
                if location:
                    title_line += f" ({location})"
                if dates:
                    title_line += f" | {dates}"
                
                edu_unit.append(Paragraph(title_line, self.styles["JobTitle"]))
                
                if gpa:
                    edu_unit.append(Paragraph(f"GPA: {gpa}", self.styles["Body"]))
                
                if honors:
                    edu_unit.append(Paragraph(f"Honors: {honors}", self.styles["Body"]))
                
                # Keep each education entry together
                education_content.append(KeepTogether(edu_unit))
                
                # Add spacing between education entries (but not after last)
                if edu != education[-1]:
                    education_content.append(Spacer(1, self.SPACE_BETWEEN_JOB_COMPONENTS))
            
            sections.append({
                "name": "EDUCATION",
                "content": education_content
            })
        
        # Additional info for abbreviated versions only (at the very end)
        if self.length_variant == "short":
            additional_info = self.data.get("additional_info", "")
            if additional_info:
                # Create clickable links for LinkedIn and Personal Site
                linkedin_url = "https://www.linkedin.com/in/dheerajchand/"
                personal_site_url = "https://www.dheerajchand.com"
                
                # Replace the entire text with URLs included
                additional_info_with_links = f"For a more detailed, full description of my experience, please visit my LinkedIn ({linkedin_url}) and Personal Site ({personal_site_url})."
                
                sections.append({
                    "name": "",
                    "content": [Paragraph(additional_info_with_links, self.styles["Body"])]
                })
        
        # Technical Skills (moved to end like Deepak) - WITH VISUAL HIERARCHY
        competencies = self.data.get("competencies", {})
        if competencies:
            technical_skills_content = []
            for main_category, sub_skills in competencies.items():
                if isinstance(sub_skills, list):
                    # Technical skills - format as Deepak does with visual hierarchy
                    colors = self.config if self.config else {}
                    accent_color = colors.get("ACCENT_COLOR", "#4682B4") 
                    muted_color = colors.get("MEDIUM_TEXT_COLOR", "#666666")
                    
                    sub_content = []
                    for skill_line in sub_skills:
                        if ": " in skill_line:
                            sub_category, details = skill_line.split(": ", 1)
                            sub_content.append(f'<font color="{accent_color}"><i>{sub_category}</i></font> <font color="{muted_color}">({details})</font>')
                        else:
                            sub_content.append(f'<font color="{accent_color}">{skill_line}</font>')
                    
                    # Create single paragraph with main category and all sub-skills using color hierarchy
                    main_color = colors.get("COMPETENCY_HEADER_COLOR", "#2C3E50")
                    skill_text = "; ".join(sub_content)
                    full_text = f'<font color="{main_color}"><b>{main_category.upper()}</b></font> {skill_text}'
                    technical_skills_content.append(Paragraph(full_text, self.styles["CompetencyDetail"]))
            
            if technical_skills_content:
                # Technical Skills - don't use KeepTogether to avoid excessive spacing
                # Individual skill paragraphs will naturally flow across pages if needed
                sections.append({
                    "name": "TECHNICAL SKILLS",
                    "content": technical_skills_content
                })
        
        return sections
    
    def _create_keep_together_section(self, content_lines, min_lines=3):
        """
        Create a KeepTogether section with configurable minimum lines
        
        NOTE: This method should only be used for short sections like:
        - Professional Summary (single paragraph)
        - Key Achievements (single line with bullets)
        - Core Competencies (single line with bullets)
        
        For longer sections like Projects, Education, Experience, use individual
        KeepTogether logic to prevent excessive spacing issues.
        
        Args:
            content_lines: List of content elements (Paragraphs, etc.)
            min_lines: Minimum number of lines to keep together (default: 3)
        
        Returns:
            KeepTogether object containing all content
        """
        if not content_lines:
            return []
        
        # Always wrap in KeepTogether for these short sections
        return [KeepTogether(content_lines)]
    
    def _create_horizontal_bar(self, color="#2C3E50", height=2):
        """Create a horizontal bar for section separation"""
        table = Table([['']], colWidths=[BAR_WIDTH], rowHeights=[height/72*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), HexColor(color)),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        return table
    
    def _calculate_header_footer_dimensions(self):
        """SYSTEMATIC APPROACH: Calculate header bar position for first page (most content)"""
        # Use the first page header content for spacing calculation since it has the most content
        # This ensures consistent spacing regardless of page type
        personal_info = self.data.get("personal_info", {})
        contact_info = personal_info.get("contact", {})
        
        # Find the lowest text line using the same logic as the header function
        # This represents the actual header bar position on the first page
        lowest_line_y = HEADER_TOP_Y  # Start with email line
        
        # Check if phone exists
        phone = contact_info.get("phone", "")
        if phone:
            phone_y = HEADER_TOP_Y - HEADER_PHONE_OFFSET
            lowest_line_y = min(lowest_line_y, phone_y)
        
        # Check if GitHub exists
        github_url = contact_info.get("github", "")
        if github_url:
            github_y = HEADER_TOP_Y - HEADER_GITHUB_OFFSET
            lowest_line_y = min(lowest_line_y, github_y)
        
        # Right side text positions (first page has all elements)
        name_y = HEADER_TOP_Y - HEADER_NAME_OFFSET
        slogan_y = name_y - 12
        austin_y = name_y - 24
        
        lowest_line_y = min(lowest_line_y, name_y, slogan_y, austin_y)
        
        # Header bar position: one line height below lowest text
        header_bar_y = lowest_line_y - FONT_SIZE_9
        
        # Footer bar position: fixed at bottom
        footer_bar_y = FOOTER_BAR_POSITION
        
        # Body content spacing: use proper spacing constant
        # This spacing will be the same for all pages since we use the same Spacer
        body_start_y = header_bar_y - SPACE_HEADER_TO_CONTENT
        body_end_y = footer_bar_y + FONT_SIZE_11
        
        # Calculate margins based on actual positions
        top_margin = PAGE_HEIGHT - body_start_y
        bottom_margin = body_end_y
        
        return {
            'header_bar_y': header_bar_y,
            'footer_bar_y': footer_bar_y,
            'body_start_y': body_start_y,
            'body_end_y': body_end_y,
            'top_margin': top_margin,
            'bottom_margin': bottom_margin
        }

    def generate_pdf(self, filename: str) -> str:
        """Generate PDF resume using systematic header/footer approach"""
        # SYSTEMATIC APPROACH: Calculate dimensions first
        dimensions = self._calculate_header_footer_dimensions()
        
        doc = SimpleDocTemplate(filename, pagesize=letter, 
                              rightMargin=MARGIN_RIGHT, leftMargin=MARGIN_LEFT,
                              topMargin=dimensions['top_margin'], 
                              bottomMargin=dimensions['bottom_margin'],
                              # Optimized quality settings for maximum quality within size limits
                              compress=1,  # Light compression for reasonable file size
                              creator="Resume Generator Pro",
                              title=f"Resume - {self.data.get('personal_info', {}).get('name', 'Professional')}",
                              author=self.data.get('personal_info', {}).get('name', 'Professional'),
                              # Additional quality settings
                              subject="Professional Resume",
                              keywords="resume, professional, career")
        story = []
        
        # Define sections in order
        sections = self._get_sections()
        
        # No additional spacing needed - topMargin already accounts for proper spacing
        # The dimensions calculation ensures consistent spacing across all pages
        
        # Render each section
        for i, section in enumerate(sections):
            if section["content"]:
                story.append(Paragraph(section["name"], self.styles["SectionHeader"]))
                story.extend(section["content"])
        
        
        
        # Build with custom header and footer using systematic approach
        def add_header(canvas, doc):
            """Add header with dynamic bar positioning for all pages"""
            canvas.saveState()
            
            # High quality rendering settings
            canvas.setLineCap(1)  # Round line caps for smoother lines
            canvas.setLineJoin(1)  # Round line joins for smoother lines
            canvas.setLineWidth(1.0)  # Ensure consistent line width
            
            personal_info = self.data.get("personal_info", {})
            contact_info = personal_info.get("contact", {})
            
            # Three-cell layout: Email/Phone (left) | Empty (middle) | Name (right)
            name = personal_info.get("name", "NAME")
            email = contact_info.get("email", "")
            phone = contact_info.get("phone", "").replace(".", "").replace("-", "").replace(" ", "")
            
            # Add US country code to phone if it doesn't have one
            if phone and not phone.startswith("+1") and not phone.startswith("1"):
                phone = f"+1 {phone}"
            elif phone and phone.startswith("1") and not phone.startswith("+1"):
                phone = f"+{phone}"
            
            # Left cell: Email and phone stacked vertically
            left_x = HEADER_LEFT_X
            top_y = HEADER_TOP_Y  # Lowered to prevent bleeding off page
            
            # Use the same logic as dimensions calculation for consistency
            # This ensures header bar position matches the spacing calculation
            lowest_line_y = top_y  # Start with email line
            
            if phone:
                phone_y = top_y - HEADER_PHONE_OFFSET
                lowest_line_y = min(lowest_line_y, phone_y)
            
            # GitHub link under phone on left side with equal spacing
            github_url = contact_info.get("github", "")
            if github_url:
                github_y = top_y - HEADER_GITHUB_OFFSET
                lowest_line_y = min(lowest_line_y, github_y)
            
            # Right side text
            name_y = top_y - HEADER_NAME_OFFSET
            slogan_y = name_y - 12
            austin_y = name_y - 24
            
            lowest_line_y = min(lowest_line_y, name_y, slogan_y, austin_y)
            
            # Get colors from config for better visual hierarchy
            primary_color = self.config.get("ACCENT_COLOR", "#4682B4")
            secondary_color = self.config.get("COMPANY_COLOR", "#2C3E50") 
            text_color = self.config.get("DARK_TEXT_COLOR", "#2C3E50")
            muted_color = self.config.get("MEDIUM_TEXT_COLOR", "#666666")
            link_color = self.config.get("LINK_COLOR", "#4682B4")
            
            # LEFT SIDE: Email and Phone (primary accent color)
            if email:
                canvas.setFont("Helvetica", FONT_SIZE_11)
                canvas.setFillColor(HexColor(primary_color))
                canvas.drawString(left_x, top_y, email)
            
            if phone:
                canvas.setFont("Helvetica", FONT_SIZE_11)
                canvas.setFillColor(HexColor(primary_color))
                canvas.drawString(left_x, phone_y, phone)
            
            # RIGHT SIDE: Name (bigger, aligned with email/phone)
            canvas.setFont("Helvetica-Bold", FONT_SIZE_14)  # Bigger name
            canvas.setFillColor(HexColor(text_color))
            canvas.drawRightString(PAGE_RIGHT_MARGIN, top_y, name)  # Aligned with email
            
            # Slogan (smaller, between name and coordinates)
            canvas.setFont("Helvetica", FONT_SIZE_9)
            canvas.setFillColor(HexColor(muted_color))
            slogan = "[RESEARCH, ANALYSIS, ENGINEERING] → UNDERSTANDING"
            canvas.drawRightString(PAGE_RIGHT_MARGIN, top_y - 14, slogan)
            
            # GitHub (aligned with coordinates, different color)
            if github_url:
                canvas.setFont("Helvetica", FONT_SIZE_9)
                # Label in secondary color
                canvas.setFillColor(HexColor(secondary_color))
                canvas.drawString(left_x, austin_y, "GitHub: ")
                github_label_width = canvas.stringWidth("GitHub: ", "Helvetica", FONT_SIZE_9)
                
                # Link in different color
                canvas.setFillColor(HexColor(link_color))
                github_text = github_url.replace('https://', '').replace('http://', '')
                github_x = left_x + github_label_width
                # Create clickable link
                canvas.linkURL(github_url, 
                             (github_x, austin_y - 2, github_x + canvas.stringWidth(github_text, "Helvetica", FONT_SIZE_9), austin_y + 10))
                canvas.drawString(github_x, austin_y, github_text)
            
            # Austin, TX with coordinates (aligned with GitHub)
            canvas.setFont("Helvetica", FONT_SIZE_9)
            canvas.setFillColor(HexColor(muted_color))
            austin_text = "Austin, TX (30.2672°N, 97.7431°W)"
            austin_width = canvas.stringWidth(austin_text, "Helvetica", FONT_SIZE_9)
            austin_x = PAGE_RIGHT_MARGIN - austin_width
            # Create clickable link to OpenStreetMap
            canvas.linkURL("https://www.openstreetmap.org/?mlat=30.2672&mlon=-97.7431&zoom=12", 
                         (austin_x, austin_y - 2, PAGE_RIGHT_MARGIN, austin_y + 10))
            canvas.drawRightString(PAGE_RIGHT_MARGIN, austin_y, austin_text)
            
            # Add horizontal bar separator - simple: one line height below lowest text
            bar_y = lowest_line_y - FONT_SIZE_9  # One line height below lowest text
            canvas.setStrokeColor(HexColor(self.config.get("MEDIUM_TEXT_COLOR", "#666666")))
            canvas.setLineWidth(BAR_LINE_WIDTH)
            canvas.line(PAGE_LEFT_MARGIN, bar_y, PAGE_RIGHT_MARGIN, bar_y)
            
            canvas.restoreState()
            add_footer(canvas, doc)
        
        
        def add_footer(canvas, doc):
            """Add footer with two-cell structure and dynamic bar positioning"""
            canvas.saveState()
            
            # High quality rendering settings
            canvas.setLineCap(1)  # Round line caps for smoother lines
            canvas.setLineJoin(1)  # Round line joins for smoother lines
            canvas.setLineWidth(1.0)  # Ensure consistent line width
            
            personal_info = self.data.get("personal_info", {})
            contact_info = personal_info.get("contact", {})
            
            canvas.setFont("Helvetica", FONT_SIZE_8)
            footer_y = FOOTER_Y
            
            # Calculate where the bar should go BEFORE drawing any text
            # Bar should be 8 points above the footer text
            bar_y = footer_y + 8
            
            # Draw the bar FIRST (above the text)
            canvas.setStrokeColor(HexColor(self.config.get("MEDIUM_TEXT_COLOR", "#666666")))
            canvas.setLineWidth(BAR_LINE_WIDTH_FOOTER)
            canvas.line(PAGE_LEFT_MARGIN, bar_y, PAGE_RIGHT_MARGIN, bar_y)
            
            # Two-cell footer structure
            website_url = contact_info.get("website", "")
            linkedin_url = contact_info.get("linkedin", "")
            
            # Left cell: Site and LinkedIn with pipe separator
            if website_url or linkedin_url:
                current_x = PAGE_LEFT_MARGIN
                
                # Draw labels in accent color, links in different color
                if website_url:
                    # Label in accent color
                    canvas.setFillColor(HexColor(self.config.get("ACCENT_COLOR", "#4682B4")))
                    canvas.drawString(current_x, footer_y, "Site: ")
                    current_x += canvas.stringWidth("Site: ", "Helvetica", FONT_SIZE_8)
                    
                    # Link in different color (use medium text color for contrast)
                    canvas.setFillColor(HexColor(self.config.get("MEDIUM_TEXT_COLOR", "#666666")))
                    canvas.linkURL(website_url, (current_x, footer_y - FOOTER_LINK_OFFSET, current_x + len(website_url)*FOOTER_LINK_OFFSET, footer_y + FOOTER_LINK_OFFSET*2))
                    canvas.drawString(current_x, footer_y, website_url)
                    current_x += canvas.stringWidth(website_url, "Helvetica", FONT_SIZE_8)
                
                if linkedin_url:
                    # Pipe separator
                    if website_url:
                        canvas.setFillColor(HexColor("#666666"))
                        canvas.drawString(current_x, footer_y, " | ")
                        current_x += canvas.stringWidth(" | ", "Helvetica", FONT_SIZE_8)
                    
                    # Label in accent color
                    canvas.setFillColor(HexColor(self.config.get("ACCENT_COLOR", "#4682B4")))
                    canvas.drawString(current_x, footer_y, "LinkedIn: ")
                    current_x += canvas.stringWidth("LinkedIn: ", "Helvetica", FONT_SIZE_8)
                    
                    # Link in different color (use medium text color for contrast)
                    canvas.setFillColor(HexColor(self.config.get("MEDIUM_TEXT_COLOR", "#666666")))
                    canvas.linkURL(linkedin_url, (current_x, footer_y - FOOTER_LINK_OFFSET, current_x + len(linkedin_url)*FOOTER_LINK_OFFSET, footer_y + FOOTER_LINK_OFFSET*2))
                    canvas.drawString(current_x, footer_y, linkedin_url)
            
            # Right cell: Page number
            canvas.setFillColor(HexColor(self.config.get("ACCENT_COLOR", "#4682B4")))
            page_num = canvas.getPageNumber()
            canvas.drawRightString(PAGE_RIGHT_MARGIN, footer_y, f"Page {page_num}")
            
            canvas.restoreState()
        
        doc.build(story, onFirstPage=add_header, onLaterPages=add_header)
        return filename
    
    def generate_docx(self, filename: str) -> str:
        """Generate DOCX resume with high quality settings"""
        doc = Document()
        
        # High quality DOCX settings
        doc.core_properties.title = f"Resume - {self.data.get('personal_info', {}).get('name', 'Professional')}"
        doc.core_properties.author = self.data.get('personal_info', {}).get('name', 'Professional')
        doc.core_properties.creator = "Resume Generator Pro"
        doc.core_properties.subject = "Professional Resume"
        doc.core_properties.keywords = "resume, professional, career"
        doc.core_properties.comments = "Generated by Resume Generator Pro"
        
        # Personal info
        personal_info = self.data.get("personal_info", {})
        contact_info = personal_info.get("contact", {})
        
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal_info.get("name", "NAME"))
        name_run.font.size = Inches(0.2)  # Keep this as it's a reasonable size for DOCX
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_parts = []
        if contact_info.get("phone"):
            contact_parts.append(contact_info["phone"])
        if contact_info.get("email"):
            contact_parts.append(contact_info["email"])
        if contact_info.get("website"):
            contact_parts.append(contact_info["website"])
        if contact_info.get("linkedin"):
            contact_parts.append(contact_info["linkedin"])
        if contact_info.get("location"):
            contact_parts.append(contact_info["location"])
        
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        summary = self.data.get("summary", "")
        if summary:
            doc.add_heading("PROFESSIONAL SUMMARY", level=2)
            doc.add_paragraph(summary)
        
        # CORE COMPETENCIES - Horizontal bullet format
        competencies = self.data.get("competencies", {})
        if competencies:
            competency_categories = []
            for main_category, sub_skills in competencies.items():
                if isinstance(sub_skills, list):
                    competency_categories.append(main_category)
            
            # Single line with bullet separators
            competency_text = " • ".join(competency_categories)
            doc.add_heading("CORE COMPETENCIES", level=2)
            doc.add_paragraph(competency_text)
        
        # Experience
        experience = self.data.get("experience", [])
        if experience:
            doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
            for job in experience:
                job_title = job.get("title", "")
                company = job.get("company", "")
                location = job.get("location", "")
                dates = job.get("dates", "")
                
                title_line = f"{job_title}"
                if company:
                    title_line += f" - {company}"
                if location:
                    title_line += f" ({location})"
                if dates:
                    title_line += f" | {dates}"
                
                doc.add_heading(title_line, level=3)
                
                if job.get("subtitle"):
                    doc.add_paragraph(job["subtitle"])
                
                responsibilities = job.get("responsibilities", [])
                for resp in responsibilities:
                    doc.add_paragraph(f"• {resp}")
        
        # Projects
        projects = self.data.get("projects", [])
        if projects:
            doc.add_heading("KEY PROJECTS", level=2)
            for project in projects:
                project_name = project.get("name", "")
                dates = project.get("dates", "")
                description = project.get("description", "")
                technologies = project.get("technologies", [])
                impact = project.get("impact", "")
                
                title_line = project_name
                if dates:
                    title_line += f" ({dates})"
                
                doc.add_heading(title_line, level=3)
                
                if description:
                    doc.add_paragraph(description)
                
                if technologies:
                    doc.add_paragraph(f"Technologies: {', '.join(technologies)}")
                
                if impact:
                    doc.add_paragraph(f"Impact: {impact}")
        
        # Education
        education = self.data.get("education", [])
        if education:
            doc.add_heading("EDUCATION", level=2)
            for edu in education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                location = edu.get("location", "")
                dates = edu.get("dates", "")
                gpa = edu.get("gpa", "")
                honors = edu.get("honors", "")
                
                title_line = degree
                if institution:
                    title_line += f" - {institution}"
                if location:
                    title_line += f" ({location})"
                if dates:
                    title_line += f" | {dates}"
                
                doc.add_heading(title_line, level=3)
                
                if gpa:
                    doc.add_paragraph(f"GPA: {gpa}")
                
                if honors:
                    doc.add_paragraph(f"Honors: {honors}")
        
        # Achievements
        achievements = self.data.get("achievements", {})
        if achievements:
            doc.add_heading("KEY ACHIEVEMENTS AND IMPACT", level=2)
            for category, achievement_list in achievements.items():
                if isinstance(achievement_list, list):
                    doc.add_heading(category, level=3)
                    for achievement in achievement_list:
                        doc.add_paragraph(f"• {achievement}")
        
        # Technical Skills (moved to end like Deepak)
        competencies = self.data.get("competencies", {})
        if competencies:
            doc.add_heading("TECHNICAL SKILLS", level=2)
            for main_category, sub_skills in competencies.items():
                if isinstance(sub_skills, list):
                    # Technical skills - format as Deepak does with specific technologies
                    skill_text = "; ".join([skill.split(": ")[0] if ": " in skill else skill for skill in sub_skills])
                    doc.add_paragraph(f"{main_category.upper()} {skill_text}")
        
        # Additional info for abbreviated versions
        additional_info = self.data.get("additional_info", "")
        if additional_info:
            doc.add_paragraph(additional_info)
        
        doc.save(filename)
        return filename
    
    def generate_rtf(self, filename: str) -> str:
        """Generate RTF resume"""
        # RTF is a text format, so we'll create a simple text version
        content = []
        
        # Personal info
        personal_info = self.data.get("personal_info", {})
        contact_info = personal_info.get("contact", {})
        content.append(f"\\b {personal_info.get('name', 'NAME')}\\b0")
        content.append("")
        
        # Contact info
        contact_parts = []
        if contact_info.get("phone"):
            contact_parts.append(contact_info["phone"])
        if contact_info.get("email"):
            contact_parts.append(contact_info["email"])
        if contact_info.get("website"):
            contact_parts.append(contact_info["website"])
        if contact_info.get("linkedin"):
            contact_parts.append(contact_info["linkedin"])
        if contact_info.get("location"):
            contact_parts.append(contact_info["location"])
        
        if contact_parts:
            content.append(" | ".join(contact_parts))
        
        content.append("")
        
        # Summary
        summary = self.data.get("summary", "")
        if summary:
            content.append("\\b PROFESSIONAL SUMMARY\\b0")
            content.append(summary)
            content.append("")
        
        # Achievements (moved to top)
        achievements = self.data.get("achievements", {})
        if achievements:
            content.append("\\b KEY ACHIEVEMENTS AND IMPACT\\b0")
            for category, achievement_list in achievements.items():
                if isinstance(achievement_list, list):
                    content.append(f"\\b {category}\\b0")
                    for achievement in achievement_list:
                        content.append(f"• {achievement}")
                    content.append("")
            content.append("")
        
        # Additional info for abbreviated versions
        additional_info = self.data.get("additional_info", "")
        if additional_info:
            if isinstance(additional_info, list):
                for info in additional_info:
                    content.append(info)
            else:
                content.append(additional_info)
            content.append("")
        
        # CORE COMPETENCIES - Space-efficient skills section
        competencies = self.data.get("competencies", {})
        if competencies:
            content.append("\\b CORE COMPETENCIES\\b0")
            for category, skills in competencies.items():
                if isinstance(skills, list):
                    content.append(f"• {category}")
            content.append("")
            
            content.append("\\b TECHNICAL SKILLS\\b0")
            for category, skills in competencies.items():
                if isinstance(skills, list):
                    if category.upper() in ["CODE", "COMPUTE", "INTERACT", "MEASURE", "PLATFORMS", "TRACK"]:
                        # Technical skills - format as Deepak does
                        skill_text = "; ".join([str(skill).split(": ")[0] if ": " in str(skill) else str(skill) for skill in skills if isinstance(skill, str)])
                        content.append(f"\\b {category.upper()}\\b0 {skill_text}")
                    else:
                        # Other skills - format as bullet points
                        for skill in skills:
                            if isinstance(skill, str):
                                skill_name = skill.split(": ")[0] if ": " in skill else skill
                                content.append(f"• {skill_name}")
            content.append("")
        
        # Experience
        experience = self.data.get("experience", [])
        if experience:
            content.append("\\b PROFESSIONAL EXPERIENCE\\b0")
            for job in experience:
                job_title = job.get("title", "")
                company = job.get("company", "")
                location = job.get("location", "")
                dates = job.get("dates", "")
                
                title_line = f"{job_title}"
                if company:
                    title_line += f" - {company}"
                if location:
                    title_line += f" ({location})"
                if dates:
                    title_line += f" | {dates}"
                
                content.append(f"\\b {title_line}\\b0")
                
                if job.get("subtitle"):
                    content.append(job["subtitle"])
                
                responsibilities = job.get("responsibilities", [])
                for resp in responsibilities:
                    content.append(f"• {resp}")
                
                content.append("")
        
        # Projects
        projects = self.data.get("projects", [])
        if projects:
            content.append("\\b KEY PROJECTS\\b0")
            for project in projects:
                project_name = project.get("name", "")
                dates = project.get("dates", "")
                description = project.get("description", "")
                technologies = project.get("technologies", [])
                impact = project.get("impact", "")
                
                title_line = project_name
                if dates:
                    title_line += f" ({dates})"
                
                content.append(f"\\b {title_line}\\b0")
                
                if description:
                    content.append(description)
                
                if technologies:
                    content.append(f"Technologies: {', '.join(technologies)}")
                
                if impact:
                    content.append(f"Impact: {impact}")
                
                content.append("")
        
        # Education
        education = self.data.get("education", [])
        if education:
            content.append("\\b EDUCATION\\b0")
            for edu in education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                location = edu.get("location", "")
                dates = edu.get("dates", "")
                gpa = edu.get("gpa", "")
                honors = edu.get("honors", "")
                
                title_line = degree
                if institution:
                    title_line += f" - {institution}"
                if location:
                    title_line += f" ({location})"
                if dates:
                    title_line += f" | {dates}"
                
                content.append(f"\\b {title_line}\\b0")
                
                if gpa:
                    content.append(f"GPA: {gpa}")
                
                if honors:
                    content.append(f"Honors: {honors}")
                
                content.append("")
        
        
        # Write RTF file
        rtf_content = "{\\rtf1\\ansi\\deff0\\par " + "\\par ".join(content) + "\\par }"
        
        with open(filename, "w", encoding="utf-8") as f:
            f.write(rtf_content)
        
        return filename
    
    def generate_markdown(self, filename: str) -> str:
        """Generate Markdown resume"""
        content = []
        
        # Personal info
        personal_info = self.data.get("personal_info", {})
        contact_info = personal_info.get("contact", {})
        content.append(f"# {personal_info.get('name', 'NAME')}")
        content.append("")
        
        # Contact information
        contact_parts = []
        if contact_info.get('phone'):
            contact_parts.append(f"**Phone:** {contact_info['phone']}")
        if contact_info.get('email'):
            contact_parts.append(f"**Email:** {contact_info['email']}")
        if contact_info.get('website'):
            contact_parts.append(f"**Website:** {contact_info['website']}")
        if contact_info.get('linkedin'):
            contact_parts.append(f"**LinkedIn:** {contact_info['linkedin']}")
        if contact_info.get('location'):
            contact_parts.append(f"**Location:** {contact_info['location']}")
        
        if contact_parts:
            content.append(" | ".join(contact_parts))
            content.append("")
        
        # Summary
        summary = self.data.get("summary", "")
        if summary:
            content.append("## Professional Summary")
            content.append("")
            content.append(summary)
            content.append("")
        
        # Achievements (moved to top)
        achievements = self.data.get("achievements", {})
        if achievements:
            content.append("## Key Achievements and Impact")
            content.append("")
            for category, achievement_list in achievements.items():
                content.append(f"### {category}")
                if isinstance(achievement_list, list):
                    for achievement in achievement_list:
                        content.append(f"- {achievement}")
                content.append("")
        
        # Additional info for abbreviated versions
        additional_info = self.data.get("additional_info", "")
        if additional_info:
            if isinstance(additional_info, list):
                for info in additional_info:
                    content.append(info)
            else:
                content.append(additional_info)
            content.append("")
        
        # CORE COMPETENCIES - Space-efficient skills section
        competencies = self.data.get("competencies", {})
        if competencies:
            content.append("## Core Competencies")
            content.append("")
            for category, skills in competencies.items():
                if isinstance(skills, list):
                    content.append(f"• **{category}**")
            content.append("")
            
            content.append("## Technical Skills")
            content.append("")
            for category, skills in competencies.items():
                if isinstance(skills, list):
                    # Show detailed skills with descriptions for all categories
                    for skill in skills:
                        if isinstance(skill, str):
                            content.append(f"• **{skill}**")
            content.append("")
        
        # Experience (enhanced formatting with visual hierarchy)
        experience = self.data.get("experience", [])
        if experience:
            content.append("## Professional Experience")
            content.append("")
            for job in experience:
                # Job title: ### header (like primary color)
                content.append(f"### {job.get('title', '')}")
                
                # Company info: **bold** (like accent color)
                company_info = job.get('company', '')
                if job.get('location'):
                    company_info += f" | {job['location']}"
                if job.get('dates'):
                    company_info += f" | {job['dates']}"
                content.append(f"**{company_info}**")
                content.append("")
                
                # Subtitle: *italics* (like muted color)
                if job.get('subtitle'):
                    content.append(f"*{job['subtitle']}*")
                    content.append("")
                
                # Responsibilities - clean formatting without tech term highlighting
                responsibilities = job.get('responsibilities', [])
                if responsibilities:
                    for resp in responsibilities:
                        content.append(f"- {resp}")
                content.append("")
        
        # Projects
        projects = self.data.get("projects", [])
        if projects:
            content.append("## Key Projects")
            content.append("")
            for project in projects:
                content.append(f"### {project.get('name', '')}")
                if project.get('dates'):
                    content.append(f"*{project['dates']}*")
                    content.append("")
                if project.get('description'):
                    content.append(f"{project['description']}")
                    content.append("")
                if project.get('technologies'):
                    content.append(f"**Technologies:** {', '.join(project['technologies'])}")
                if project.get('impact'):
                    content.append(f"**Impact:** {project['impact']}")
                content.append("")
        
        # Education
        education = self.data.get("education", [])
        if education:
            content.append("## Education")
            content.append("")
            for edu in education:
                content.append(f"### {edu.get('degree', '')}")
                institution_info = edu.get('institution', '')
                if edu.get('location'):
                    institution_info += f" | {edu['location']}"
                if edu.get('dates'):
                    institution_info += f" | {edu['dates']}"
                content.append(f"**{institution_info}**")
                if edu.get('gpa'):
                    content.append(f"**GPA:** {edu['gpa']}")
                if edu.get('honors'):
                    content.append(f"**Honors:** {edu['honors']}")
                content.append("")
        
        
        # Footer with contact information
        content.append("---")
        content.append("")
        
        # Add footer contact info similar to PDF footer
        personal_info = self.data.get("personal_info", {})
        contact_info = personal_info.get("contact", {})
        footer_parts = []
        
        if contact_info.get("website"):
            footer_parts.append(f"**Website:** {contact_info['website']}")
        if contact_info.get("linkedin"):
            footer_parts.append(f"**LinkedIn:** {contact_info['linkedin']}")
            
        if footer_parts:
            content.append(" | ".join(footer_parts))
        
        with open(filename, "w", encoding="utf-8") as f:
            f.write("\n".join(content))
        
        return filename


class ResumeManager:
    """Manages resume generation with color schemes and formats"""
    
    def __init__(self):
        self.versions = {
            "comprehensive": "dheeraj_chand_comprehensive_full",
            "polling_research_redistricting": "dheeraj_chand_polling_research_redistricting",
            "marketing": "dheeraj_chand_marketing",
            "data_analysis_visualization": "dheeraj_chand_data_analysis_visualization",
            "data_engineering": "dheeraj_chand_data_engineering",
            "product": "dheeraj_chand_product",
            "software_engineering": "dheeraj_chand_software_engineering",
            "gis": "dheeraj_chand_gis"
        }
        
        # Length variants for each version
        self.length_variants = {
            "long": "full",
            "short": "abbreviated"
        }
        
        self.color_schemes = [
            "default_professional",
            "corporate_blue", 
            "modern_tech",
            "modern_clean",
            "satellite_imagery",
            "terrain_mapping",
            "cartographic_professional",
            "topographic_classic"
        ]
        
        self.formats = ["pdf", "docx", "rtf", "md"]
    
    def generate_single_resume(self, version: str, color_scheme: str, format_type: str, output_dir: str = "outputs", length_variant: str = "long", output_type: str = "ats") -> bool:
        """Generate a single resume with specified parameters"""
        if version not in self.versions:
            return False
        
        if length_variant not in self.length_variants:
            return False
        
        # Determine input directory based on length variant and output type
        input_basename = self.versions[version]
        if output_type == "human":
            input_basename += "_human"
        if length_variant == "short":
            input_basename += "_abbreviated"
        
        input_dir = Path("inputs") / input_basename
        
        if not input_dir.exists():
            return False
        
        data_file = input_dir / "resume_data.json"
        config_file = input_dir / "config.json"
        
        if not data_file.exists():
            return False
        
        try:
            # Load color scheme
            color_scheme_file = Path("color_schemes") / f"{color_scheme}.json"
            if color_scheme_file.exists():
                generator = ResumeGenerator(str(data_file), str(color_scheme_file), color_scheme, length_variant, output_type)
            else:
                generator = ResumeGenerator(str(data_file), str(config_file) if config_file.exists() else None, color_scheme, length_variant, output_type)
            
            # Create output directory using the correct structure: output_type/version/length/color_scheme
            base_output_dir = Path(output_dir) / output_type
            output_path = base_output_dir / version / length_variant / color_scheme
            output_path.mkdir(parents=True, exist_ok=True)
            
            # Generate file
            filename = output_path / f"dheeraj_chand_{version}_{length_variant}_{color_scheme}.{format_type}"
            
            if format_type == "pdf":
                generator.generate_pdf(str(filename))
            elif format_type == "docx":
                generator.generate_docx(str(filename))
            elif format_type == "rtf":
                generator.generate_rtf(str(filename))
            elif format_type == "md":
                generator.generate_markdown(str(filename))
            else:
                return False
            
            return True
            
        except Exception as e:
            print(f"Error generating {version} {color_scheme} {format_type}: {e}")
            return False
    
    def generate_all_combinations(self, output_dir: str = "outputs", output_type: str = "ats") -> Dict[str, int]:
        """Generate all combinations of versions, lengths, color schemes, and formats"""
        results = {"success": 0, "failed": 0}
        
        for version in self.versions:
            for length_variant in self.length_variants:
                for color_scheme in self.color_schemes:
                    for format_type in self.formats:
                        if self.generate_single_resume(version, color_scheme, format_type, output_dir, length_variant, output_type):
                            results["success"] += 1
                        else:
                            results["failed"] += 1
        
        return results
